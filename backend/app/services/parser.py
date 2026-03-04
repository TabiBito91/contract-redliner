"""Document parsing service - extracts structured content from DOCX files.

This module implements the document parsing abstraction layer (NFR-09) so that
new formats (PDF in Phase 2) can be added without modifying the comparison engine.
"""

import re
from abc import ABC, abstractmethod
from pathlib import Path
from uuid import UUID

from docx import Document as DocxDocument
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from app.models.schemas import (
    DocumentParagraph,
    DocumentTable,
    ParsedDocument,
)


class DocumentParser(ABC):
    """Abstract base class for document parsers (NFR-09 extensibility)."""

    @abstractmethod
    def parse(self, file_path: Path, document_id: UUID) -> ParsedDocument:
        """Parse a document file into the intermediate representation."""
        ...

    @abstractmethod
    def supports(self, file_path: Path) -> bool:
        """Check if this parser supports the given file type."""
        ...


class DocxParser(DocumentParser):
    """Parser for .docx files using python-docx."""

    # Regex to detect section/clause numbering patterns
    _section_number_re = re.compile(
        r"^(\d+(?:\.\d+)*\.?)\s"  # e.g., "1.", "1.1", "8.2.1"
    )

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".docx"

    def parse(self, file_path: Path, document_id: UUID) -> ParsedDocument:
        doc = DocxDocument(str(file_path))

        paragraphs = self._extract_paragraphs(doc)
        tables = self._extract_tables(doc)
        headers = self._extract_headers(doc)
        footers = self._extract_footers(doc)
        footnotes = self._extract_footnotes(doc)

        return ParsedDocument(
            document_id=document_id,
            paragraphs=paragraphs,
            tables=tables,
            headers=headers,
            footers=footers,
            footnotes=footnotes,
        )

    def _extract_paragraphs(self, doc: DocxDocument) -> list[DocumentParagraph]:
        """Extract all body paragraphs with structure metadata."""
        paragraphs = []
        current_heading_section = None

        for i, para in enumerate(doc.paragraphs):
            text = self._normalize_text(para.text)
            if not text:
                continue

            heading_level = self._get_heading_level(para)
            section_number = self._extract_section_number(text)
            list_level, list_marker = self._get_list_info(para)

            if heading_level is not None:
                current_heading_section = section_number or text[:50]

            paragraphs.append(
                DocumentParagraph(
                    id=f"p-{i}",
                    text=text,
                    heading_level=heading_level,
                    section_number=section_number,
                    list_level=list_level,
                    list_marker=list_marker,
                    style_name=para.style.name if para.style else None,
                    parent_section=current_heading_section if heading_level is None else None,
                )
            )

        return paragraphs

    def _extract_tables(self, doc: DocxDocument) -> list[DocumentTable]:
        """Extract tables with cell content."""
        tables = []
        for i, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [self._normalize_text(cell.text) for cell in row.cells]
                rows.append(cells)
            tables.append(
                DocumentTable(
                    id=f"t-{i}",
                    rows=rows,
                )
            )
        return tables

    def _extract_headers(self, doc: DocxDocument) -> list[DocumentParagraph]:
        """Extract header content from all sections."""
        headers = []
        for i, section in enumerate(doc.sections):
            header = section.header
            if header and not header.is_linked_to_previous:
                for j, para in enumerate(header.paragraphs):
                    text = self._normalize_text(para.text)
                    if text:
                        headers.append(
                            DocumentParagraph(
                                id=f"hdr-{i}-{j}",
                                text=text,
                            )
                        )
        return headers

    def _extract_footers(self, doc: DocxDocument) -> list[DocumentParagraph]:
        """Extract footer content from all sections."""
        footers = []
        for i, section in enumerate(doc.sections):
            footer = section.footer
            if footer and not footer.is_linked_to_previous:
                for j, para in enumerate(footer.paragraphs):
                    text = self._normalize_text(para.text)
                    if text:
                        footers.append(
                            DocumentParagraph(
                                id=f"ftr-{i}-{j}",
                                text=text,
                            )
                        )
        return footers

    def _extract_footnotes(self, doc: DocxDocument) -> list[DocumentParagraph]:
        """Extract footnotes and endnotes."""
        footnotes = []
        # python-docx doesn't have direct footnote access; we parse the XML
        try:
            footnotes_part = doc.part.rels.get(RT.FOOTNOTES)
            if footnotes_part:
                # Basic extraction from footnotes XML
                from lxml import etree
                nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                root = etree.fromstring(footnotes_part.target_part.blob)
                for i, fn in enumerate(root.findall(".//w:footnote", nsmap)):
                    fn_type = fn.get(f"{{{nsmap['w']}}}type")
                    if fn_type in ("separator", "continuationSeparator"):
                        continue
                    texts = fn.findall(".//w:t", nsmap)
                    text = "".join(t.text or "" for t in texts)
                    text = self._normalize_text(text)
                    if text:
                        footnotes.append(
                            DocumentParagraph(id=f"fn-{i}", text=text)
                        )
        except Exception:
            pass  # footnotes are optional; fail gracefully
        return footnotes

    def _normalize_text(self, text: str) -> str:
        """Normalize whitespace and formatting artifacts (FR-204)."""
        # Collapse multiple whitespace to single space
        text = re.sub(r"[ \t]+", " ", text)
        # Strip leading/trailing whitespace
        text = text.strip()
        # Normalize unicode quotes and dashes
        text = text.replace("\u2018", "'").replace("\u2019", "'")
        text = text.replace("\u201c", '"').replace("\u201d", '"')
        text = text.replace("\u2013", "-").replace("\u2014", "-")
        return text

    def _get_heading_level(self, para) -> int | None:
        """Detect heading level from paragraph style."""
        if para.style and para.style.name:
            name = para.style.name.lower()
            if name.startswith("heading"):
                try:
                    return int(name.replace("heading", "").strip())
                except ValueError:
                    return None
        return None

    def _extract_section_number(self, text: str) -> str | None:
        """Extract section/clause number from text start."""
        match = self._section_number_re.match(text)
        if match:
            return match.group(1).rstrip(".")
        return None

    def _get_list_info(self, para) -> tuple[int | None, str | None]:
        """Detect list level and marker from paragraph."""
        if para.style and para.style.name:
            name = para.style.name.lower()
            if "list" in name:
                # Attempt to detect nesting level from indentation
                indent = para.paragraph_format.left_indent
                level = 0
                if indent:
                    level = min(int(indent.inches / 0.25), 9)
                return level, None  # marker detection would need numPr XML parsing
        return None, None


# ---------------------------------------------------------------------------
# Shared PDF utilities — used by both PdfParser and the export generator so
# that paragraph text is always extracted and normalised identically.
# ---------------------------------------------------------------------------

def _normalize_pdf_text(text: str) -> str:
    """Normalise PDF-extracted text: remove artifacts, expand ligatures."""
    text = text.replace("\u00ad", "")                          # soft hyphen
    text = re.sub(r"[ \t\xa0\u2009\u200a]+", " ", text).strip()  # whitespace variants
    text = (text                                                # common ligatures
            .replace("\ufb00", "ff")
            .replace("\ufb01", "fi")
            .replace("\ufb02", "fl")
            .replace("\ufb03", "ffi")
            .replace("\ufb04", "ffl"))
    text = text.replace("\u2018", "'").replace("\u2019", "'")  # quotes
    text = text.replace("\u201c", '"').replace("\u201d", '"')
    text = text.replace("\u2013", "-").replace("\u2014", "-")  # dashes
    return text


def _pdf_extract_blocks(pdf_path: Path) -> list[tuple[str, float, bool]]:
    """Extract block-level text from a PDF with font metadata.

    Each PyMuPDF block corresponds to one visual paragraph.  Lines within
    a block are joined into a single string and normalised with
    _normalize_pdf_text, so consumers always see the same text regardless
    of whether they are parsing for diffing or building an export document.

    Returns: list of (block_text, max_font_size, is_bold)
    """
    import fitz  # PyMuPDF — lazy import

    blocks: list[tuple[str, float, bool]] = []
    pdf = fitz.open(str(pdf_path))
    try:
        for page in pdf:
            for block in page.get_text("dict")["blocks"]:
                if block.get("type") != 0:  # skip image blocks
                    continue

                parts: list[str] = []
                sizes: list[float] = []
                bold = False

                for line in block["lines"]:
                    spans = line.get("spans", [])
                    if not spans:
                        continue
                    line_text = _normalize_pdf_text("".join(s["text"] for s in spans))
                    if not line_text:
                        continue
                    parts.append(line_text)
                    sizes.extend(s["size"] for s in spans)
                    if any(s["flags"] & 16 for s in spans):  # bold flag = 16
                        bold = True

                if not parts:
                    continue

                block_text = _normalize_pdf_text(" ".join(parts))
                if block_text:
                    blocks.append((block_text, max(sizes), bold))
    finally:
        pdf.close()

    return blocks


class PdfParser(DocumentParser):
    """Parser for .pdf files using PyMuPDF (fitz).

    Text extraction is delegated to the shared _pdf_extract_blocks utility
    so that the export generator always sees identical paragraph text.
    """

    _section_number_re = re.compile(r"^(\d+(?:\.\d+)*\.?)\s")

    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() == ".pdf"

    def parse(self, file_path: Path, document_id: UUID) -> ParsedDocument:
        raw_blocks = _pdf_extract_blocks(file_path)
        paragraphs = self._classify_paragraphs(raw_blocks)
        return ParsedDocument(
            document_id=document_id,
            paragraphs=paragraphs,
            tables=[],
            headers=[],
            footers=[],
            footnotes=[],
        )

    def _classify_paragraphs(
        self, raw_blocks: list[tuple[str, float, bool]]
    ) -> list[DocumentParagraph]:
        """Classify pre-extracted blocks into headings and body paragraphs."""
        if not raw_blocks:
            return []

        # Body font size = median (filters large headings + small footnotes)
        all_sizes = sorted(size for _, size, _ in raw_blocks)
        body_size = all_sizes[len(all_sizes) // 2]

        paragraphs: list[DocumentParagraph] = []
        current_section: str | None = None

        for idx, (text, size, bold) in enumerate(raw_blocks):
            has_num = bool(self._section_number_re.match(text))
            is_heading = (size > body_size * 1.12 or bold) and (has_num or len(text) < 80)
            section_num = self._extract_section_number(text)

            if is_heading:
                depth = len(section_num.split(".")) if section_num else 1
                current_section = section_num or text[:50]
                paragraphs.append(DocumentParagraph(
                    id=f"p-{idx}",
                    text=text,
                    heading_level=min(depth, 9),
                    section_number=section_num,
                    style_name="Heading",
                    parent_section=None,
                ))
            else:
                paragraphs.append(DocumentParagraph(
                    id=f"p-{idx}",
                    text=text,
                    heading_level=None,
                    section_number=section_num,
                    style_name=None,
                    parent_section=current_section,
                ))

        return paragraphs

    def _extract_section_number(self, text: str) -> str | None:
        m = self._section_number_re.match(text)
        return m.group(1).rstrip(".") if m else None


def _patch_converted_docx(docx_path: Path) -> None:
    """Fix heading structure in a pdf2docx-converted DOCX.

    pdf2docx renders section headings as bold Normal-style paragraphs with a
    soft line break (\\n in para.text) separating the heading text from the
    first sentence of body text.  This function:
      1. Detects such merged paragraphs (bold + section-number prefix + \\n).
      2. Inserts a proper Heading-styled paragraph for the heading text.
      3. Rewrites the original paragraph to contain only the body text.

    Without this fix DocxParser sees no headings and treats the entire document
    as one flat provision, producing whole-section ADDITION/DELETION diffs
    instead of inline changes.
    """
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    _section_re = re.compile(r"^\d+(?:\.\d+)*\.?\s")

    doc = Document(str(docx_path))
    modified = False

    for para in list(doc.paragraphs):
        text = para.text
        if "\n" not in text:
            continue

        lines = [ln for ln in text.split("\n") if ln.strip()]
        if len(lines) < 2:
            continue

        first_line = lines[0].strip()
        if not _section_re.match(first_line):
            continue

        is_bold = any(run.bold for run in para.runs if run.text.strip())
        if not is_bold:
            continue

        # Determine heading depth from the section number (e.g. "3.1" → depth 2)
        m = re.match(r"^(\d+(?:\.\d+)*)", first_line)
        depth = len(m.group(1).split(".")) if m else 1
        level = min(depth, 6)

        # Add a proper Heading paragraph via the python-docx API (handles style
        # lookup correctly), then move it immediately before the current paragraph.
        heading_para = doc.add_heading(first_line, level=level)
        para._element.addprevious(heading_para._element)

        # Rewrite the current paragraph to contain only the body text.
        p_elem = para._element
        for child in list(p_elem):
            if child.tag in (qn("w:r"), qn("w:hyperlink")):
                p_elem.remove(child)

        # Drop any existing pStyle so the paragraph reverts to Normal.
        pPr = p_elem.find(qn("w:pPr"))
        if pPr is not None:
            pStyle = pPr.find(qn("w:pStyle"))
            if pStyle is not None:
                pPr.remove(pStyle)

        body_text = " ".join(ln.strip() for ln in lines[1:])
        body_r = OxmlElement("w:r")
        body_t = OxmlElement("w:t")
        body_t.text = body_text
        body_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        body_r.append(body_t)
        p_elem.append(body_r)

        modified = True

    if modified:
        doc.save(str(docx_path))


def convert_pdf_to_docx(pdf_path: Path) -> Path:
    """Convert a PDF to a temporary DOCX file using pdf2docx.

    The caller is responsible for deleting the returned file when done.
    Uses a UUID-suffixed name in the system temp directory to avoid collisions
    when multiple comparisons run concurrently.
    """
    import os
    import tempfile
    from pdf2docx import Converter

    fd, tmp_str = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    tmp_path = Path(tmp_str)
    cv = Converter(str(pdf_path))
    try:
        cv.convert(str(tmp_path), start=0, end=None)
    finally:
        cv.close()
    return tmp_path


class ParserRegistry:
    """Registry of available document parsers (supports Phase 2 extensibility)."""

    def __init__(self):
        self._parsers: list[DocumentParser] = []

    def register(self, parser: DocumentParser) -> None:
        self._parsers.append(parser)

    def get_parser(self, file_path: Path) -> DocumentParser:
        for parser in self._parsers:
            if parser.supports(file_path):
                return parser
        raise ValueError(f"No parser available for file type: {file_path.suffix}")


# Global parser registry
parser_registry = ParserRegistry()
parser_registry.register(DocxParser())
parser_registry.register(PdfParser())
