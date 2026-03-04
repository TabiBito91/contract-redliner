"""Output document generation service (ADR-004: Clone-and-Annotate).

Clones the original DOCX and inserts redline markup:
- Deletions: red text with strikethrough
- Additions: blue text with underline
- AI annotations: Word comments at change locations
- Summary appendix: table at end of document
"""

import copy
from pathlib import Path
from uuid import UUID

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

from app.models.schemas import (
    AnnotatedChange,
    ChangeType,
    ExportOptions,
    RiskSeverity,
)
from app.services.diff_engine import DiffChange, InlineDiff


# --- Color constants ---
RED = RGBColor(0xCC, 0x00, 0x00)       # Deletions
BLUE = RGBColor(0x00, 0x44, 0xCC)      # Additions
ORANGE = RGBColor(0xFF, 0x88, 0x00)    # Moves
GRAY = RGBColor(0x88, 0x88, 0x88)      # Cosmetic/info

SEVERITY_COLORS = {
    RiskSeverity.CRITICAL: RGBColor(0xCC, 0x00, 0x00),
    RiskSeverity.HIGH: RGBColor(0xFF, 0x66, 0x00),
    RiskSeverity.MEDIUM: RGBColor(0xDD, 0xAA, 0x00),
    RiskSeverity.LOW: RGBColor(0x00, 0x88, 0x00),
    RiskSeverity.INFO: RGBColor(0x00, 0x44, 0xCC),
}


def _add_formatted_run(paragraph, text: str, color: RGBColor,
                       strikethrough: bool = False, underline: bool = False,
                       bold: bool = False):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.color.rgb = color
    if strikethrough:
        run.font.strike = True
    if underline:
        run.font.underline = True
    if bold:
        run.bold = True
    return run


def _add_comment(doc, paragraph, comment_text: str, author: str = "RedlineAI"):
    """Add a Word comment to a paragraph.

    This uses low-level OOXML manipulation since python-docx doesn't
    natively support comments.
    """
    # Get or create the comments part
    comments_part = None
    for rel in doc.part.rels.values():
        if "comments" in rel.reltype:
            comments_part = rel.target_part
            break

    if comments_part is None:
        # Create a comments part - this is complex OOXML; skip for MVP
        # Comments will be implemented as inline annotations instead
        return

    # For MVP, we'll use a simpler approach: add the comment as
    # a bracketed annotation at the end of the paragraph
    run = paragraph.add_run(f"  [{comment_text}]")
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    run.font.italic = True


def _normalize_ws(text: str) -> str:
    """Normalize whitespace for matching."""
    import re
    return re.sub(r"\s+", " ", text).strip()


def _extract_section_number(text: str) -> float | None:
    """Return the leading section number from a heading string, or None.

    Handles formats like '3.', '3.1', '10.' at the start of the text.
    """
    import re
    if not text:
        return None
    m = re.match(r"^(\d+(?:\.\d+)?)", text.strip())
    return float(m.group(1)) if m else None


def _build_section_end_map(doc_paragraphs) -> dict[float, int]:
    """Map each section number to the index of its last paragraph.

    Scans the document for lines that start with a section number and
    tracks the furthest paragraph index seen under each section heading.
    """
    section_map: dict[float, int] = {}
    current_section: float | None = None
    for i, para in enumerate(doc_paragraphs):
        num = _extract_section_number(para.text)
        if num is not None:
            current_section = num
        if current_section is not None:
            section_map[current_section] = i  # advances to last para of section
    return section_map


def _match_para_index(
    doc_paragraphs, original_text: str, threshold: float = 0.8
) -> int | None:
    """Find the best matching paragraph index in the DOCX by text content.

    threshold: minimum SequenceMatcher ratio to accept a match.  Use 0.65
    for PDF-derived documents where minor encoding differences are expected.
    """
    if not original_text:
        return None
    norm_orig = _normalize_ws(original_text)
    if not norm_orig:
        return None

    best_idx = None
    best_ratio = 0.0
    from difflib import SequenceMatcher
    for i, para in enumerate(doc_paragraphs):
        para_text = _normalize_ws(para.text)
        if not para_text:
            continue
        # Quick exact check
        if para_text == norm_orig:
            return i
        # Fuzzy match for minor whitespace/unicode differences
        ratio = SequenceMatcher(None, para_text, norm_orig).ratio()
        if ratio > best_ratio:
            best_ratio = ratio
            best_idx = i

    return best_idx if best_ratio > threshold else None


def _clear_paragraph(para):
    """Remove all runs from a paragraph, keeping the paragraph element."""
    for run in para.runs:
        para._element.remove(run._element)
    # Also clear any direct text nodes
    p_elem = para._element
    for child in list(p_elem):
        if child.tag.endswith("}r"):
            p_elem.remove(child)


def _apply_inline_redline(para, original_text: str, modified_text: str):
    """Replace a paragraph's content with inline redline markup."""
    _clear_paragraph(para)

    from difflib import SequenceMatcher
    matcher = SequenceMatcher(None, original_text or "", modified_text or "", autojunk=False)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            para.add_run(original_text[i1:i2])
        elif tag == "delete":
            _add_formatted_run(para, original_text[i1:i2], RED, strikethrough=True)
        elif tag == "insert":
            _add_formatted_run(para, modified_text[j1:j2], BLUE, underline=True)
        elif tag == "replace":
            _add_formatted_run(para, original_text[i1:i2], RED, strikethrough=True)
            _add_formatted_run(para, modified_text[j1:j2], BLUE, underline=True)


def _build_doc_from_pdf(original_path: Path) -> Document:
    """Convert a PDF to a plain python-docx Document for redline export.

    Uses the shared _pdf_extract_blocks utility (same as PdfParser) so that
    paragraph texts match annotated_changes.original_text exactly, allowing
    _match_para_index to find deletions and modifications reliably.
    """
    from app.services.parser import _pdf_extract_blocks

    doc = Document()
    for text, _, _ in _pdf_extract_blocks(original_path):
        doc.add_paragraph(text)
    return doc


def generate_redline_docx(
    original_path: Path,
    annotated_changes: list[AnnotatedChange],
    options: ExportOptions,
    output_path: Path,
) -> Path:
    """Generate a redlined DOCX document.

    For DOCX originals, clones the file and annotates it in place.
    For PDF originals, extracts the text into a fresh DOCX first, then
    applies the same annotation logic.

    Args:
        original_path: Path to the original document (DOCX or PDF).
        annotated_changes: Changes with optional AI summaries and risk assessments.
        options: Export configuration toggles.
        output_path: Where to save the output DOCX.

    Returns:
        Path to the generated output file.
    """
    if original_path.suffix.lower() == ".pdf":
        doc = _build_doc_from_pdf(original_path)
        match_threshold = 0.65  # P3: PDF text may differ slightly from parsed text
    else:
        doc = Document(str(original_path))
        match_threshold = 0.8

    # Filter changes based on options
    active_changes: list[AnnotatedChange] = []
    for ac in annotated_changes:
        if not ac.change.is_substantive and not options.show_formatting_changes:
            continue
        active_changes.append(ac)

    doc_paragraphs = list(doc.paragraphs)
    used_indices: set[int] = set()

    # Build section-number → last-paragraph-index map for anchor lookups
    section_end_map = _build_section_end_map(doc_paragraphs)

    # Apply modifications and deletions to matched paragraphs
    for ac in active_changes:
        ct = ac.change.change_type
        if ct == ChangeType.ADDITION:
            continue

        idx = _match_para_index(doc_paragraphs, ac.change.original_text, match_threshold)
        if idx is None or idx in used_indices:
            continue
        used_indices.add(idx)

        if ct == ChangeType.DELETION:
            original_text = doc_paragraphs[idx].text
            _clear_paragraph(doc_paragraphs[idx])
            _add_formatted_run(doc_paragraphs[idx], original_text, RED, strikethrough=True)
        else:
            # MODIFICATION, MOVE, FORMAT_ONLY
            _apply_inline_redline(
                doc_paragraphs[idx],
                ac.change.original_text or "",
                ac.change.modified_text or "",
            )
        if options.include_ai_summaries and ac.ai_summary:
            _add_comment(doc, doc_paragraphs[idx], ac.ai_summary.summary)

    # Insert additions at their correct document position using section numbering.
    # insertion_tracker maps anchor_paragraph_idx → the last element inserted
    # after it, so consecutive additions chain correctly (2 → A → B → C).
    from docx.text.paragraph import Paragraph as DocxParagraph
    insertion_tracker: dict[int, object] = {}

    for ac in active_changes:
        if ac.change.change_type != ChangeType.ADDITION:
            continue

        # Determine which section number this addition belongs to
        add_section_num = _extract_section_number(
            ac.change.section_context or ac.change.modified_text or ""
        )

        # Find the last paragraph of the nearest preceding section
        anchor_idx: int | None = None
        if add_section_num is not None and section_end_map:
            preceding = [s for s in section_end_map if s < add_section_num]
            if preceding:
                anchor_idx = section_end_map[max(preceding)]

        # Determine the actual element to insert after
        if anchor_idx is not None:
            insert_after_elem = insertion_tracker.get(anchor_idx) or doc_paragraphs[anchor_idx]._element
        else:
            insert_after_elem = None

        # Create and insert the new paragraph
        new_p = OxmlElement("w:p")
        if insert_after_elem is not None:
            insert_after_elem.addnext(new_p)
        else:
            doc.element.body.append(new_p)

        para = DocxParagraph(new_p, doc._body)
        _add_formatted_run(para, "[ADDED] ", BLUE, bold=True)
        _add_formatted_run(para, ac.change.modified_text or "", BLUE, underline=True)
        if options.include_ai_summaries and ac.ai_summary:
            _add_comment(doc, para, ac.ai_summary.summary)

        # Update tracker so the next addition after the same anchor chains on
        if anchor_idx is not None:
            insertion_tracker[anchor_idx] = new_p

    # Insert disclaimer header at top
    disclaimer_para = doc.add_paragraph()
    disclaimer_para.paragraph_format.space_after = Pt(12)
    _add_formatted_run(
        disclaimer_para,
        "REDLINE COMPARISON - Generated by RedlineAI | ",
        GRAY, bold=True,
    )
    _add_formatted_run(
        disclaimer_para,
        "AI assessments do not constitute legal advice.",
        GRAY,
    )
    doc.element.body.insert(0, disclaimer_para._element)

    # Add legend
    legend_para = doc.add_paragraph()
    _add_formatted_run(legend_para, "Deleted text", RED, strikethrough=True)
    _add_formatted_run(legend_para, "  |  ", GRAY)
    _add_formatted_run(legend_para, "Added text", BLUE, underline=True)
    _add_formatted_run(legend_para, "  |  ", GRAY)
    _add_formatted_run(legend_para, "Moved text", ORANGE, underline=True)
    doc.element.body.insert(1, legend_para._element)

    # Add summary appendix if enabled
    if options.include_summary_appendix:
        _add_summary_appendix(doc, annotated_changes, options)

    doc.save(str(output_path))
    return output_path


def _add_summary_appendix(
    doc: Document,
    annotated_changes: list[AnnotatedChange],
    options: ExportOptions,
):
    """Append a summary table at the end of the document."""
    doc.add_page_break()
    doc.add_heading("Change Summary", level=1)

    # Disclaimer
    disclaimer = doc.add_paragraph()
    _add_formatted_run(
        disclaimer,
        "AI-generated summaries and risk assessments. This does not constitute legal advice.",
        GRAY, bold=True,
    )
    disclaimer.paragraph_format.space_after = Pt(12)

    # Determine columns
    columns = ["#", "Type", "Section"]
    if options.include_ai_summaries:
        columns.append("Summary")
    if options.include_risk_assessments:
        columns.extend(["Risk", "Recommendation"])

    # Create table
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"

    # Header row
    for i, col_name in enumerate(columns):
        cell = table.rows[0].cells[i]
        cell.text = col_name
        for run in cell.paragraphs[0].runs:
            run.bold = True

    # Data rows
    substantive_idx = 0
    for ac in annotated_changes:
        change = ac.change
        if not change.is_substantive and not options.show_formatting_changes:
            continue

        substantive_idx += 1
        row = table.add_row()
        col = 0

        row.cells[col].text = str(substantive_idx)
        col += 1

        row.cells[col].text = change.change_type.value.replace("_", " ").title()
        col += 1

        row.cells[col].text = change.section_context or ""
        col += 1

        if options.include_ai_summaries:
            summary_text = ac.ai_summary.summary if ac.ai_summary else ""
            row.cells[col].text = summary_text
            col += 1

        if options.include_risk_assessments:
            if ac.risk_assessment:
                severity = ac.risk_assessment.severity.value.upper()
                row.cells[col].text = severity
                # Color the risk cell
                para = row.cells[col].paragraphs[0]
                if para.runs:
                    color = SEVERITY_COLORS.get(ac.risk_assessment.severity, GRAY)
                    para.runs[0].font.color.rgb = color
                    para.runs[0].bold = True
                col += 1

                row.cells[col].text = ac.risk_assessment.recommendation
                col += 1
            else:
                row.cells[col].text = ""
                col += 1
                row.cells[col].text = ""
                col += 1


def generate_inline_redline_paragraph(
    original_text: str,
    modified_text: str,
    inline_diffs: list[InlineDiff],
) -> list[tuple[str, str]]:
    """Generate a sequence of (text, format) tuples for inline redline display.

    Returns list of tuples: (text_content, format_type)
    where format_type is one of: "normal", "deletion", "addition"
    """
    if not original_text and not modified_text:
        return [("No changes detected", "normal")]

    if not original_text:
        return [(modified_text, "addition")]
    if not modified_text:
        return [(original_text, "deletion")]

    # Reconstruct the full inline view by diffing the texts directly
    from difflib import SequenceMatcher
    matcher = SequenceMatcher(None, original_text, modified_text, autojunk=False)

    result = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            result.append((original_text[i1:i2], "normal"))
        elif tag == "delete":
            result.append((original_text[i1:i2], "deletion"))
        elif tag == "insert":
            result.append((modified_text[j1:j2], "addition"))
        elif tag == "replace":
            result.append((original_text[i1:i2], "deletion"))
            result.append((modified_text[j1:j2], "addition"))

    return result
