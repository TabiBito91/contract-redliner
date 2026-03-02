"""Generate sample DOCX contracts for diff engine testing.

Creates an Original and a Modified version with known changes:
- Modified definitions
- Deleted provisions
- Added new provisions
- Moved/reordered clauses
- Minor word substitutions
- Formatting-only changes
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt

FIXTURES = Path(__file__).parent / "fixtures"
FIXTURES.mkdir(exist_ok=True)


def make_original():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    doc.add_heading("MUTUAL NON-DISCLOSURE AGREEMENT", level=1)

    doc.add_heading("1. Definitions", level=2)
    doc.add_paragraph(
        '1.1 "Confidential Information" means any and all non-public, proprietary, '
        "or trade secret information disclosed by either party to the other party, "
        "whether in written, oral, electronic, or other form, that is designated as "
        '"confidential" or that reasonably should be understood to be confidential '
        "given the nature of the information and the circumstances of disclosure."
    )
    doc.add_paragraph(
        '1.2 "Disclosing Party" means the party disclosing Confidential Information.'
    )
    doc.add_paragraph(
        '1.3 "Receiving Party" means the party receiving Confidential Information.'
    )
    doc.add_paragraph(
        '1.4 "Representatives" means a party\'s employees, officers, directors, '
        "advisors, agents, and consultants who have a need to know the Confidential "
        "Information for the Purpose."
    )

    doc.add_heading("2. Purpose", level=2)
    doc.add_paragraph(
        "2.1 The parties wish to explore a potential business relationship "
        '(the "Purpose") and, in connection therewith, each party may disclose '
        "Confidential Information to the other party."
    )

    doc.add_heading("3. Obligations of the Receiving Party", level=2)
    doc.add_paragraph(
        "3.1 The Receiving Party shall hold and maintain the Confidential Information "
        "in strict confidence and shall not, without the prior written consent of the "
        "Disclosing Party, disclose any Confidential Information to any third party."
    )
    doc.add_paragraph(
        "3.2 The Receiving Party shall use the Confidential Information solely for "
        "the Purpose and shall not use the Confidential Information for any other "
        "purpose without the prior written consent of the Disclosing Party."
    )
    doc.add_paragraph(
        "3.3 The Receiving Party shall limit disclosure of Confidential Information "
        "to its Representatives who have a need to know and who are bound by "
        "confidentiality obligations no less restrictive than those contained herein."
    )

    doc.add_heading("4. Exclusions from Confidential Information", level=2)
    doc.add_paragraph(
        "4.1 Confidential Information shall not include information that: "
        "(a) is or becomes publicly available through no fault of the Receiving Party; "
        "(b) was known to the Receiving Party prior to disclosure by the Disclosing Party; "
        "(c) is independently developed by the Receiving Party without use of the "
        "Confidential Information; or "
        "(d) is rightfully received from a third party without restriction on disclosure."
    )

    doc.add_heading("5. Term and Termination", level=2)
    doc.add_paragraph(
        "5.1 This Agreement shall remain in effect for a period of two (2) years "
        'from the date of execution (the "Term").'
    )
    doc.add_paragraph(
        "5.2 Either party may terminate this Agreement at any time by providing "
        "thirty (30) days prior written notice to the other party."
    )
    doc.add_paragraph(
        "5.3 The obligations of confidentiality shall survive the termination of "
        "this Agreement for a period of three (3) years."
    )

    doc.add_heading("6. Return of Materials", level=2)
    doc.add_paragraph(
        "6.1 Upon termination of this Agreement or upon request by the Disclosing "
        "Party, the Receiving Party shall promptly return or destroy all copies of "
        "the Confidential Information in its possession."
    )

    doc.add_heading("7. Remedies", level=2)
    doc.add_paragraph(
        "7.1 The parties acknowledge that any breach of this Agreement may cause "
        "irreparable harm to the Disclosing Party for which monetary damages would "
        "be inadequate. Accordingly, the Disclosing Party shall be entitled to seek "
        "equitable relief, including injunction and specific performance, in addition "
        "to all other remedies available at law or in equity."
    )

    doc.add_heading("8. Limitation of Liability", level=2)
    doc.add_paragraph(
        "8.1 In no event shall either party be liable to the other party for any "
        "indirect, incidental, special, consequential, or punitive damages arising "
        "out of or related to this Agreement, regardless of whether such damages are "
        "based on contract, tort, strict liability, or any other theory."
    )
    doc.add_paragraph(
        "8.2 The total aggregate liability of either party under this Agreement "
        "shall not exceed Five Million Dollars ($5,000,000)."
    )

    doc.add_heading("9. Governing Law", level=2)
    doc.add_paragraph(
        "9.1 This Agreement shall be governed by and construed in accordance with "
        "the laws of the State of New York, without regard to its conflict of laws "
        "principles."
    )

    doc.add_heading("10. General Provisions", level=2)
    doc.add_paragraph(
        "10.1 This Agreement constitutes the entire agreement between the parties "
        "with respect to the subject matter hereof and supersedes all prior and "
        "contemporaneous agreements and understandings."
    )
    doc.add_paragraph(
        "10.2 This Agreement may not be amended except by a written instrument "
        "signed by both parties."
    )
    doc.add_paragraph(
        "10.3 Neither party may assign this Agreement without the prior written "
        "consent of the other party."
    )

    path = FIXTURES / "nda_original.docx"
    doc.save(str(path))
    print(f"Created: {path}")
    return path


def make_modified():
    """Create a modified version with specific, trackable changes."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    doc.add_heading("MUTUAL NON-DISCLOSURE AGREEMENT", level=1)

    # --- CHANGE 1: Modified definition (narrowed scope) ---
    doc.add_heading("1. Definitions", level=2)
    doc.add_paragraph(
        '1.1 "Confidential Information" means any and all non-public, proprietary '
        "information disclosed by either party to the other party, "
        "whether in written or electronic form, that is explicitly designated as "
        '"confidential" at the time of disclosure.'
        # REMOVED: "or trade secret" from scope
        # REMOVED: "oral" as a form
        # REMOVED: "reasonably should be understood" standard
        # CHANGED: requires explicit designation at time of disclosure
    )
    doc.add_paragraph(
        '1.2 "Disclosing Party" means the party disclosing Confidential Information.'
    )
    doc.add_paragraph(
        '1.3 "Receiving Party" means the party receiving Confidential Information.'
    )
    doc.add_paragraph(
        # --- CHANGE 2: Modified definition of Representatives (added contractors) ---
        '1.4 "Representatives" means a party\'s employees, officers, directors, '
        "advisors, agents, consultants, and independent contractors who have a need "
        "to know the Confidential Information for the Purpose."
    )

    doc.add_heading("2. Purpose", level=2)
    doc.add_paragraph(
        "2.1 The parties wish to explore a potential business relationship "
        '(the "Purpose") and, in connection therewith, each party may disclose '
        "Confidential Information to the other party."
    )

    doc.add_heading("3. Obligations of the Receiving Party", level=2)
    doc.add_paragraph(
        # --- CHANGE 3: Weakened from "strict confidence" to "reasonable confidence" ---
        "3.1 The Receiving Party shall hold and maintain the Confidential Information "
        "in reasonable confidence and shall not, without the prior written consent of the "
        "Disclosing Party, disclose any Confidential Information to any third party."
    )
    doc.add_paragraph(
        "3.2 The Receiving Party shall use the Confidential Information solely for "
        "the Purpose and shall not use the Confidential Information for any other "
        "purpose without the prior written consent of the Disclosing Party."
    )
    doc.add_paragraph(
        "3.3 The Receiving Party shall limit disclosure of Confidential Information "
        "to its Representatives who have a need to know and who are bound by "
        "confidentiality obligations no less restrictive than those contained herein."
    )

    doc.add_heading("4. Exclusions from Confidential Information", level=2)
    doc.add_paragraph(
        "4.1 Confidential Information shall not include information that: "
        "(a) is or becomes publicly available through no fault of the Receiving Party; "
        "(b) was known to the Receiving Party prior to disclosure by the Disclosing Party; "
        "(c) is independently developed by the Receiving Party without use of the "
        "Confidential Information; or "
        "(d) is rightfully received from a third party without restriction on disclosure."
    )

    # --- CHANGE 4: MOVED Section 9 (Governing Law) up before Term ---
    doc.add_heading("5. Governing Law", level=2)
    doc.add_paragraph(
        "5.1 This Agreement shall be governed by and construed in accordance with "
        "the laws of the State of Delaware, without regard to its conflict of laws "
        # --- CHANGE 5: Changed from New York to Delaware ---
        "principles."
    )

    doc.add_heading("6. Term and Termination", level=2)
    doc.add_paragraph(
        # --- CHANGE 6: Changed term from 2 years to 3 years ---
        "6.1 This Agreement shall remain in effect for a period of three (3) years "
        'from the date of execution (the "Term").'
    )
    doc.add_paragraph(
        # --- CHANGE 7: Changed notice period from 30 to 15 days ---
        "6.2 Either party may terminate this Agreement at any time by providing "
        "fifteen (15) days prior written notice to the other party."
    )
    doc.add_paragraph(
        # --- CHANGE 8: Changed survival from 3 years to 5 years ---
        "6.3 The obligations of confidentiality shall survive the termination of "
        "this Agreement for a period of five (5) years."
    )

    doc.add_heading("7. Return of Materials", level=2)
    doc.add_paragraph(
        "7.1 Upon termination of this Agreement or upon request by the Disclosing "
        "Party, the Receiving Party shall promptly return or destroy all copies of "
        "the Confidential Information in its possession."
    )

    # --- CHANGE 9: Section 7 (Remedies) DELETED entirely ---
    # (Remedies section removed)

    doc.add_heading("8. Limitation of Liability", level=2)
    doc.add_paragraph(
        "8.1 In no event shall either party be liable to the other party for any "
        "indirect, incidental, special, consequential, or punitive damages arising "
        "out of or related to this Agreement, regardless of whether such damages are "
        "based on contract, tort, strict liability, or any other theory."
    )
    doc.add_paragraph(
        # --- CHANGE 10: Reduced liability cap from $5M to $2M ---
        "8.2 The total aggregate liability of either party under this Agreement "
        "shall not exceed Two Million Dollars ($2,000,000)."
    )

    # --- CHANGE 11: NEW section added ---
    doc.add_heading("9. Non-Solicitation", level=2)
    doc.add_paragraph(
        "9.1 During the Term and for a period of one (1) year thereafter, neither "
        "party shall directly or indirectly solicit or hire any employee of the other "
        "party who was involved in the exchange of Confidential Information."
    )

    doc.add_heading("10. General Provisions", level=2)
    doc.add_paragraph(
        "10.1 This Agreement constitutes the entire agreement between the parties "
        "with respect to the subject matter hereof and supersedes all prior and "
        "contemporaneous agreements and understandings."
    )
    doc.add_paragraph(
        "10.2 This Agreement may not be amended except by a written instrument "
        "signed by both parties."
    )
    doc.add_paragraph(
        "10.3 Neither party may assign this Agreement without the prior written "
        "consent of the other party."
    )

    path = FIXTURES / "nda_modified.docx"
    doc.save(str(path))
    print(f"Created: {path}")
    return path


if __name__ == "__main__":
    make_original()
    make_modified()
    print("\nKnown changes:")
    print("  1. Definition of 'Confidential Information' narrowed (removed trade secret, oral, reasonably-understood)")
    print("  2. Definition of 'Representatives' expanded (added independent contractors)")
    print("  3. 'strict confidence' weakened to 'reasonable confidence'")
    print("  4. Governing Law section MOVED from Section 9 to Section 5")
    print("  5. Governing law changed from New York to Delaware")
    print("  6. Term extended from 2 years to 3 years")
    print("  7. Notice period reduced from 30 days to 15 days")
    print("  8. Survival period extended from 3 years to 5 years")
    print("  9. Remedies section DELETED entirely")
    print(" 10. Liability cap reduced from $5M to $2M")
    print(" 11. Non-Solicitation section ADDED (entirely new)")
