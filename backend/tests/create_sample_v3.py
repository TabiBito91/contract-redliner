"""Generate a 3rd version of the NDA for multi-document testing.

V3 is based on V2 (the modified version) with additional changes:
- Counterparty pushes back on some V2 changes
- Adds new provisions
- Accepts some V2 changes as-is
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt

FIXTURES = Path(__file__).parent / "fixtures"


def make_v3():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    doc.add_heading("MUTUAL NON-DISCLOSURE AGREEMENT", level=1)

    # V3 CHANGE: Partially reverts V2's narrowing - adds back "trade secret" but keeps
    # the "explicitly designated" requirement
    doc.add_heading("1. Definitions", level=2)
    doc.add_paragraph(
        '1.1 "Confidential Information" means any and all non-public, proprietary, '
        "or trade secret information disclosed by either party to the other party, "
        "whether in written or electronic form, that is explicitly designated as "
        '"confidential" at the time of disclosure.'
    )
    doc.add_paragraph(
        '1.2 "Disclosing Party" means the party disclosing Confidential Information.'
    )
    doc.add_paragraph(
        '1.3 "Receiving Party" means the party receiving Confidential Information.'
    )
    # Keeps V2's addition of independent contractors
    doc.add_paragraph(
        '1.4 "Representatives" means a party\'s employees, officers, directors, '
        "advisors, agents, consultants, and independent contractors who have a need "
        "to know the Confidential Information for the Purpose."
    )

    doc.add_heading("2. Purpose", level=2)
    doc.add_paragraph(
        "2.1 The parties wish to explore a potential business relationship "
        '(the "Purpose") and, in connection therewith, each party may disclose '
        "Confidential Information to the other party."
    )

    doc.add_heading("3. Obligations of the Receiving Party", level=2)
    # V3 CHANGE: Compromises between original "strict" and V2's "reasonable" -> "commercially reasonable"
    doc.add_paragraph(
        "3.1 The Receiving Party shall hold and maintain the Confidential Information "
        "using commercially reasonable measures of confidence and shall not, without the prior "
        "written consent of the Disclosing Party, disclose any Confidential Information "
        "to any third party."
    )
    doc.add_paragraph(
        "3.2 The Receiving Party shall use the Confidential Information solely for "
        "the Purpose and shall not use the Confidential Information for any other "
        "purpose without the prior written consent of the Disclosing Party."
    )
    doc.add_paragraph(
        "3.3 The Receiving Party shall limit disclosure of Confidential Information "
        "to its Representatives who have a need to know and who are bound by "
        "confidentiality obligations no less restrictive than those contained herein."
    )

    doc.add_heading("4. Exclusions from Confidential Information", level=2)
    doc.add_paragraph(
        "4.1 Confidential Information shall not include information that: "
        "(a) is or becomes publicly available through no fault of the Receiving Party; "
        "(b) was known to the Receiving Party prior to disclosure by the Disclosing Party; "
        "(c) is independently developed by the Receiving Party without use of the "
        "Confidential Information; or "
        "(d) is rightfully received from a third party without restriction on disclosure."
    )

    # V3 keeps V2's move of Governing Law to Section 5
    doc.add_heading("5. Governing Law", level=2)
    doc.add_paragraph(
        "5.1 This Agreement shall be governed by and construed in accordance with "
        "the laws of the State of Delaware, without regard to its conflict of laws "
        "principles."
    )

    doc.add_heading("6. Term and Termination", level=2)
    # V3 CHANGE: Compromises on term - 3 years from V2 accepted
    doc.add_paragraph(
        "6.1 This Agreement shall remain in effect for a period of three (3) years "
        'from the date of execution (the "Term").'
    )
    # V3 CHANGE: Pushes back on V2's 15 days -> compromises to 20 days
    doc.add_paragraph(
        "6.2 Either party may terminate this Agreement at any time by providing "
        "twenty (20) days prior written notice to the other party."
    )
    # V3 CHANGE: Pushes back on V2's 5 years survival -> compromises to 4 years
    doc.add_paragraph(
        "6.3 The obligations of confidentiality shall survive the termination of "
        "this Agreement for a period of four (4) years."
    )

    doc.add_heading("7. Return of Materials", level=2)
    doc.add_paragraph(
        "7.1 Upon termination of this Agreement or upon request by the Disclosing "
        "Party, the Receiving Party shall promptly return or destroy all copies of "
        "the Confidential Information in its possession."
    )

    # V3 CHANGE: Remedies section RESTORED (was deleted in V2)
    doc.add_heading("8. Remedies", level=2)
    doc.add_paragraph(
        "8.1 The parties acknowledge that any breach of this Agreement may cause "
        "irreparable harm to the Disclosing Party for which monetary damages would "
        "be inadequate. Accordingly, the Disclosing Party shall be entitled to seek "
        "equitable relief, including injunction and specific performance, in addition "
        "to all other remedies available at law or in equity."
    )

    doc.add_heading("9. Limitation of Liability", level=2)
    doc.add_paragraph(
        "9.1 In no event shall either party be liable to the other party for any "
        "indirect, incidental, special, consequential, or punitive damages arising "
        "out of or related to this Agreement, regardless of whether such damages are "
        "based on contract, tort, strict liability, or any other theory."
    )
    # V3 CHANGE: Compromises on cap - $3M instead of V2's $2M (original was $5M)
    doc.add_paragraph(
        "9.2 The total aggregate liability of either party under this Agreement "
        "shall not exceed Three Million Dollars ($3,000,000)."
    )

    # V3 CHANGE: Keeps V2's non-solicitation but narrows scope
    doc.add_heading("10. Non-Solicitation", level=2)
    doc.add_paragraph(
        "10.1 During the Term and for a period of six (6) months thereafter, neither "
        "party shall directly solicit any employee of the other "
        "party who was involved in the exchange of Confidential Information."
    )

    # V3 CHANGE: NEW section - Dispute Resolution (added by V3)
    doc.add_heading("11. Dispute Resolution", level=2)
    doc.add_paragraph(
        "11.1 Any dispute arising out of or relating to this Agreement shall first "
        "be submitted to non-binding mediation. If mediation is unsuccessful, the "
        "dispute shall be resolved by binding arbitration in accordance with the "
        "rules of the American Arbitration Association."
    )

    doc.add_heading("12. General Provisions", level=2)
    doc.add_paragraph(
        "12.1 This Agreement constitutes the entire agreement between the parties "
        "with respect to the subject matter hereof and supersedes all prior and "
        "contemporaneous agreements and understandings."
    )
    doc.add_paragraph(
        "12.2 This Agreement may not be amended except by a written instrument "
        "signed by both parties."
    )
    doc.add_paragraph(
        "12.3 Neither party may assign this Agreement without the prior written "
        "consent of the other party."
    )

    path = FIXTURES / "nda_v3.docx"
    doc.save(str(path))
    print(f"Created: {path}")
    return path


if __name__ == "__main__":
    make_v3()
    print("\nV3 changes from V2:")
    print("  1. Confidential Info: partially reverts (adds back 'trade secret')")
    print("  2. Confidence standard: 'reasonable' -> 'commercially reasonable measures'")
    print("  3. Notice period: 15 days -> 20 days (compromise)")
    print("  4. Survival period: 5 years -> 4 years (compromise)")
    print("  5. Remedies section RESTORED")
    print("  6. Liability cap: $2M -> $3M (compromise)")
    print("  7. Non-solicitation: narrowed (removed 'indirectly', 1yr->6mo)")
    print("  8. Dispute Resolution: NEW section added")
